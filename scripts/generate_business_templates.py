from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('apps/web/public/business-templates')
OUT.mkdir(parents=True, exist_ok=True)
GREEN = '0B5D3B'; GOLD = 'C9A43B'; LIGHT = 'EAF3EE'; GRAY = 'F2F4F3'; INK = '24332C'

def font(run, size=10.5, bold=False, color=INK):
    run.font.name = 'Heiti SC'
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'Heiti SC')
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc.get_or_add_tcPr(); tcMar = tc.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tc.append(tcMar)
    for side, val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = OxmlElement(f'w:{side}'); node.set(qn('w:w'), str(val)); node.set(qn('w:type'),'dxa'); tcMar.append(node)

def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr(); rep = OxmlElement('w:tblHeader'); rep.set(qn('w:val'),'true'); trPr.append(rep)

def base_doc(title, code, version='V1.0'):
    d = Document(); s = d.sections[0]
    s.page_width=Inches(8.5); s.page_height=Inches(11); s.top_margin=Inches(.72); s.bottom_margin=Inches(.72); s.left_margin=Inches(.72); s.right_margin=Inches(.72)
    styles=d.styles
    normal=styles['Normal']; normal.font.name='Heiti SC'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Heiti SC'); normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(INK)
    normal.paragraph_format.space_after=Pt(5); normal.paragraph_format.line_spacing=1.15
    for name,size,color in [('Title',22,GREEN),('Heading 1',15,GREEN),('Heading 2',12,GREEN),('Heading 3',11,'4E6B5D')]:
        st=styles[name]; st.font.name='Heiti SC'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Heiti SC'); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=True
    header=s.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run(f'金特夫企业经营管理系统  |  {code}'),8.5,False,'61736A')
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(3)
    font(p.add_run('金 特 夫  KING TURF'),12,True,GREEN)
    p=d.add_paragraph(); p.style='Title'; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run(title)
    meta=d.add_table(rows=2, cols=4); meta.alignment=WD_TABLE_ALIGNMENT.CENTER; meta.autofit=False
    set_repeat_header(meta.rows[0])
    labels=[('文件编号',code),('版本',version),('生效日期','〔填写〕'),('受控状态','受控模板')]
    for i,(k,v) in enumerate(labels):
        r=i//2; c=(i%2)*2
        meta.columns[c].width=Inches(1.0); meta.columns[c+1].width=Inches(2.35)
        meta.cell(r,c).text=k; meta.cell(r,c+1).text=v; shade(meta.cell(r,c),LIGHT)
        for cell in (meta.cell(r,c),meta.cell(r,c+1)):
            margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs: font(run,9.5,k==run.text,GREEN if k==run.text else INK)
    d.add_paragraph()
    footer=s.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    font(footer.add_run('本模板由金特夫业务文档库统一维护；使用前应核对项目、价格、税率、交期和签署权限。'),8,False,'61736A')
    return d

def table(d, headers, rows, widths=None):
    t=d.add_table(rows=1, cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    set_repeat_header(t.rows[0])
    for i,h in enumerate(headers):
        t.cell(0,i).text=h; shade(t.cell(0,i),GREEN)
        for r in t.cell(0,i).paragraphs[0].runs: font(r,9,True,'FFFFFF')
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val); margins(cells[i])
            for r in cells[i].paragraphs[0].runs: font(r,9)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Inches(w)
    return t

def h(d, text, level=1): d.add_heading(text, level=level)
def p(d, text, bold_prefix=None):
    para=d.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        font(para.add_run(bold_prefix),10.5,True,GREEN); font(para.add_run(text[len(bold_prefix):]))
    else: font(para.add_run(text))
    return para
def numbered(d, items):
    for item in items:
        para=d.add_paragraph(style='List Number'); font(para.add_run(item))
def signatures(d, parties=('甲方（盖章）：','乙方（盖章）：')):
    d.add_paragraph(); t=table(d,list(parties),[['法定代表人/授权代表：','法定代表人/授权代表：'],['签署日期：____年__月__日','签署日期：____年__月__日']], [3.35,3.35]); return t

def technical_requirement():
    d=base_doc('人造草坪项目技术需求确认书','KT-TEC-F01')
    h(d,'一、项目与使用场景'); table(d,['项目','填写内容'],[
        ['客户/项目名称','〔填写〕'],['项目地址与气候','〔填写；含高温、低温、强紫外、沿海盐雾等〕'],['用途','□十一人制足球  □五/七人制足球  □门球  □景观  □幼儿园  □其他'],['场地面积','〔填写〕㎡；净铺装面积〔填写〕㎡'],['场地基础','□水泥  □沥青  □碎石  □其他；平整度及排水情况〔填写〕'],['计划交付/施工期','〔填写〕'],['适用标准','GB/T 20394-2019；中小学项目另核对 GB 36246-2018、GB/T 43566-2023及项目所在地要求']],[1.55,5.15])
    h(d,'二、草坪系统技术指标'); table(d,['指标','确认值/选项','验收方法或说明'],[
        ['产品类型','□直丝  □曲丝  □直曲混织  □免填充','按批准样品及封样'],['草丝材质','□PE  □PP  □PA  □复合','材料声明/检测报告'],['草高','50 mm（默认足球系统，可改）','标称值允许偏差按适用标准及供方检验规范'],['磅重（Dtex）','12,000 Dtex（默认，可按方案修改）','产品规格书/检测报告'],['密度','10,500 针/㎡及以上（默认）','实测针距、簇数换算'],['针距/行距','5/8 inch（默认）','卷材实测'],['底布','PP复合底布；背胶：丁苯/PU〔选定〕','来料/成品检验'],['卷宽与卷长','4 m × 项目排版长度','排版图及卷号清单'],['颜色与线条','主色〔填写〕；划线色〔填写〕','色卡/封样'],['填充系统','石英砂〔填写〕kg/㎡；弹性颗粒〔填写〕kg/㎡','施工记录及抽检'],['渗水性能','按 GB/T 20394-2019 及批准方案','第三方或出厂检验'],['耐候/力学/有害物质','符合合同约定适用标准','具备资质的检测报告']],[1.35,2.25,3.1])
    h(d,'三、包装、标识与追溯'); p(d,'每卷须标识合同号、订单号、产品型号、规格、卷号、批次、长度、生产日期和检验状态；同一场地关键可视区域原则上使用同一色批。随货提供装箱单、出厂检验报告、合格证和维护说明。')
    h(d,'四、需求冻结与变更'); numbered(d,['本确认书经销售、技术和客户确认后作为报价、合同、采购、生产及验收的输入。','涉及草高、Dtex、密度、颜色、底布、背胶、卷宽、数量、标准或交期的变更，必须形成书面变更单并重新评审成本、库存和交期。','口头沟通、聊天记录中的意向不得替代经双方确认的版本。'])
    signatures(d,('客户确认：','金特夫技术/销售确认：')); return d

def product_spec():
    d=base_doc('体育用人造草产品技术规格书','KT-PRD-S01')
    h(d,'一、标准产品定义'); table(d,['项目','标准配置'],[
        ['产品型号','KT-FB50-12000D-165'],['推荐用途','专业训练、校园及社会足球场地（项目适用性须技术评审）'],['草丝结构','PE单丝，直丝；绿色双色'],['草高','50 mm'],['总Dtex','12,000 Dtex'],['行距','5/8 inch'],['簇密度','≥10,500 针/㎡'],['针数','165 针/米（标称）'],['底布','双层PP基布+网格增强层'],['背胶','丁苯乳胶；如需PU背胶须在订单中明确'],['标准卷宽','4.00 m'],['标准包装','卷芯+防潮外包装；单卷唯一卷号'],['执行标准','GB/T 20394-2019；项目另有更高标准时按合同技术附件执行']],[1.8,4.9])
    h(d,'二、批次检验与放行'); table(d,['检验类别','检验项目','放行要求'],[
        ['首件','颜色、草高、针距、密度、背胶、幅宽','与批准样及生产指令一致'],['过程','克重/纱线用量、簇绒外观、漏针、色差','超限隔离并评审'],['成品','外观、尺寸、卷号、标签、包装','检验合格后方可入库'],['型式/第三方','耐候、力学、有害物质、渗水等','按适用标准、客户或招标要求']],[1.2,2.5,3.0])
    h(d,'三、储运与施工提示'); numbered(d,['卷材应平码、避雨、避晒、远离热源和化学污染，装卸不得折弯、拖拽或刺破包装。','铺装前核对卷号、色批和排版图；同一区域按排版顺序展开并静置释放应力。','施工胶粘剂、接缝带、填充物和施工参数应与批准的系统方案配套，不得擅自替换。','最终产品参数以订单、批准样、批次检验报告和合同技术附件共同确定。'])
    h(d,'四、订单参数确认'); table(d,['订单号','项目','颜色/色批','卷号范围','数量（㎡）','技术确认'],[['〔填写〕','〔填写〕','〔填写〕','〔填写〕','〔填写〕','〔签字〕']],[1.0,1.25,1.05,1.3,1.0,1.1]); return d

def rfq():
    d=base_doc('采购询价及比价文件','KT-PUR-F01')
    h(d,'一、询价信息'); table(d,['项目','内容'],[['询价单号','〔填写〕'],['项目/订单','〔填写〕'],['报价截止','〔填写〕'],['交货地点','〔填写〕'],['结算币种/税率','人民币；增值税率〔填写〕%'],['联系人','〔填写〕']],[1.5,5.2])
    h(d,'二、采购清单与技术要求'); table(d,['序号','物料编码','名称/规格','单位','数量','质量/标准','要求交期','含税单价'],[
        ['1','RM-YARN-〔填写〕','PE人造草丝，Dtex/颜色/性能按技术附件','kg','〔填写〕','COA+批次追溯','〔填写〕','〔填写〕'],['2','RM-BACK-〔填写〕','PP复合底布，幅宽/克重按图纸','㎡','〔填写〕','尺寸、克重、拉力','〔填写〕','〔填写〕'],['3','RM-GLUE-〔填写〕','丁苯乳胶/PU背胶〔选定〕','kg','〔填写〕','批次检测及保质期','〔填写〕','〔填写〕']],[.45,.9,1.75,.45,.6,1.1,.75,.7])
    h(d,'三、供应商必须响应事项'); table(d,['响应项','供应商填写'],[
        ['未税/含税总价、税率及发票类型','〔填写〕'],['最小起订量、包装规格、净重/毛重','〔填写〕'],['承诺交期、产能及分批交货方案','〔填写〕'],['付款条件及账期','〔填写〕'],['质量保证期、退换货和索赔机制','〔填写〕'],['样品、检测报告、材料声明及追溯资料','□随报价提供  □可在〔日期〕前提供'],['报价有效期','自报价日起〔填写〕日'],['偏离说明','无偏离/详见附件〔填写〕']],[2.3,4.4])
    h(d,'四、内部比价与定标'); table(d,['供应商','价格30%','质量25%','交期20%','付款15%','服务10%','总分','结论'],[['A','', '', '', '', '', '', ''],['B','','','','','','',''],['C','','','','','','','']],[1.1,.75,.75,.75,.75,.75,.65,1.2])
    p(d,'定标原则：不得仅按最低价定标。采购应核对合格供应商状态、样品/检测、历史质量、产能、交期、付款条件和单一来源风险，并记录未选择最低价时的理由。'); return d

def contract_doc(title, code, purchase=False):
    d=base_doc(title,code)
    buyer='甲方（买方）' if purchase else '甲方（买方/客户）'; seller='乙方（卖方/供应商）' if purchase else '乙方（卖方）：金特夫〔完整主体名称〕'
    h(d,'合同主体与基础信息'); table(d,['项目','甲方','乙方'],[
        ['名称',buyer,seller],['统一社会信用代码','〔填写〕','〔填写〕'],['地址/联系人/电话','〔填写〕','〔填写〕'],['开户行及账号','〔填写〕','〔填写〕'],['合同编号/关联订单','〔填写〕','〔填写〕']],[1.35,2.67,2.67])
    h(d,'第一条 标的、数量与价款'); table(d,['序号','产品/物料编码','名称及规格','单位','数量','含税单价','含税金额','税率'],[['1','〔填写〕','详见本合同技术附件及批准样','〔填写〕','〔填写〕','〔填写〕','〔填写〕','〔填写〕%'],['合计','','','','','','人民币（大写）〔填写〕；￥〔填写〕','']],[.45,.9,1.55,.45,.6,.75,1.4,.6])
    clauses=[
        ('第二条 技术标准与文件优先顺序','标的物应符合本合同、技术附件、经确认的需求书/规格书、封样、图纸、订单及适用国家或行业标准。文件冲突时，双方书面确认的最新技术附件优先；任何降低质量或改变关键参数的替代均须事先书面批准。'),
        ('第三条 交付、包装与风险转移','交付地点、收货人、运输方式和交期以本合同为准。包装应满足防潮、防污、防折损和可追溯要求。风险在货物运抵约定地点并完成数量及外观签收后转移；签收不代表质量最终验收。'),
        ('第四条 验收与异议','收货方在到货后〔5〕个工作日内完成数量、包装和外观初验；隐蔽缺陷、性能或有害物质问题以检验、施工或合理使用中发现的时间起算。验收依据包括批次报告、封样、规格书和适用标准。'),
        ('第五条 价款、发票与付款','付款节点：预付款〔填写〕%；到货/验收款〔填写〕%；质保金〔填写〕%。收款方应按约定开具合法有效发票。付款不构成对质量、数量或违约责任的放弃。'),
        ('第六条 质量保证与追溯','供方保证产品为全新合格品，批次、原料、生产日期、检验、卷号/批号可追溯。质保期为验收合格后〔填写〕个月；适用法律、标准或承诺规定更长期限的，从其规定。'),
        ('第七条 变更管理','数量、规格、材料、颜色、包装、交期、地点和价格变更须由双方授权人员书面确认。未经确认的口头通知不改变合同。变更导致的成本和交期影响应在执行前明确。'),
        ('第八条 违约责任','逾期交付按逾期部分价款每日〔0.05〕%支付违约金，累计不超过合同总价〔10〕%；逾期超过〔15〕日或严重质量不合格，守约方可解除相关订单并要求赔偿可证明的损失。违约金不足以弥补损失的，违约方继续赔偿差额。'),
        ('第九条 知识产权、保密与合规','双方对在履约中知悉的价格、配方、工艺、客户、图纸和经营信息承担保密义务。供方保证交付物不侵犯第三方权利，并遵守反商业贿赂、产品质量、环境、安全和进出口合规要求。'),
        ('第十条 不可抗力','受影响方应及时通知并在合理期限内提供证明，采取措施减少损失。不可抗力持续超过〔30〕日，双方应协商变更或解除未履行部分。'),
        ('第十一条 争议解决与通知','本合同适用中华人民共和国法律。争议先协商；协商不成，向〔乙方所在地/约定地〕有管辖权的人民法院起诉。合同列明的地址、电子邮箱为有效送达地址，变更应书面通知。'),
        ('第十二条 生效、文本与附件','本合同自双方盖章或经双方认可的可靠电子签名完成之日起生效。扫描件、可靠电子签名文本与纸质盖章文本具有同等效力。附件、订单、技术确认书、封样记录和经确认的变更单均为合同组成部分。')]
    for title_,body in clauses: h(d,title_,2); p(d,body)
    h(d,'附件清单',1); table(d,['附件','名称','是否齐备'],[['1','技术需求确认书/产品规格书','□'],['2','报价单及订单清单','□'],['3','质量标准、封样及检测要求','□'],['4','交付排版/包装/标识要求','□']],[.6,5.3,.8])
    signatures(d); return d

def quality_acceptance():
    d=base_doc('人造草坪成品检验与项目验收单','KT-QUA-F01')
    h(d,'一、批次与项目识别'); table(d,['字段','内容'],[['项目/客户','〔填写〕'],['合同/订单','〔填写〕'],['产品型号/规格','〔填写〕'],['生产批次/卷号','〔填写〕'],['检验日期/地点','〔填写〕'],['抽样方案','〔填写〕']],[1.6,5.1])
    h(d,'二、检验记录'); table(d,['序号','检验项目','合同/标准要求','实测结果','判定','证据编号'],[
        ['1','外观、颜色、漏针、破损','与封样一致，无影响使用缺陷','','□合格 □不合格',''],['2','草高/幅宽/卷长','按规格书及允许偏差','','□合格 □不合格',''],['3','Dtex/密度/针距','按规格书/批次报告','','□合格 □不合格',''],['4','底布与背胶','结构、附着及外观符合要求','','□合格 □不合格',''],['5','包装、标签、卷号追溯','完整、清晰、与清单一致','','□合格 □不合格',''],['6','性能及有害物质','按适用标准/第三方报告','','□合格 □不合格','']],[.45,1.15,1.8,1.15,1.0,1.15])
    h(d,'三、不合格处置与结论'); table(d,['项目','填写内容'],[['不合格描述/位置/数量','〔填写；无则填“无”〕'],['处置','□接收  □返工  □挑选  □退货  □让步接收（需授权）'],['整改期限与责任人','〔填写〕'],['最终结论','□验收合格  □有条件接收  □验收不合格'],['附件','照片〔 〕张；检测报告编号〔填写〕；NCR编号〔填写〕']],[1.8,4.9])
    signatures(d,('客户/收货方：','检验/交付方：')); return d

def delivery():
    d=base_doc('发货、到货与签收确认单','KT-LOG-F01')
    h(d,'一、发运信息'); table(d,['字段','内容'],[['销售订单/合同','〔填写〕'],['客户/项目/收货地址','〔填写〕'],['承运商/司机/电话','〔填写〕'],['车牌/运单号','〔填写〕'],['发货时间/预计到达','〔填写〕'],['装车照片/封签号','〔填写〕']],[1.65,5.05])
    h(d,'二、随货清单'); table(d,['序号','产品型号','卷号/批号','规格','数量','包装状态','备注'],[['1','〔填写〕','〔填写〕','〔填写〕','〔填写〕','完好/异常',''],['2','','','','','','']],[.45,1.15,1.25,1.25,.7,1.0,.9])
    h(d,'三、到货核验'); table(d,['检查项','结果','异常说明/证据'],[['车辆/封签与运单一致','□是 □否',''],['包装无破损、受潮、污染','□是 □否',''],['卷号、批次、型号与清单一致','□是 □否',''],['数量与尺寸初验一致','□是 □否',''],['随货文件齐全','□是 □否',''],['异常照片及通知时间','〔填写〕','']],[2.25,1.3,3.15])
    p(d,'签收说明：本单确认到货数量、包装和可见状态，不替代合同约定的性能、隐蔽缺陷及最终质量验收。异常应当场拍照、在运单和本单中备注，并于〔24〕小时内通知交付联系人。')
    signatures(d,('收货方签收：','承运/交付方：')); return d

docs={
 '01-人造草坪项目技术需求确认书.docx':technical_requirement(),
 '02-体育用人造草产品技术规格书.docx':product_spec(),
 '03-采购询价及比价文件.docx':rfq(),
 '04-采购合同模板.docx':contract_doc('人造草坪及原辅材料采购合同','KT-PUR-C01',True),
 '05-销售合同模板.docx':contract_doc('人造草坪产品销售合同','KT-SAL-C01',False),
 '06-成品检验与项目验收单.docx':quality_acceptance(),
 '07-发货到货与签收确认单.docx':delivery(),
}
for name,doc in docs.items(): doc.save(OUT/name)
print('\n'.join(str(OUT/name) for name in docs))
